from typing import Dict, Any, List, Optional, Union
import os
import time
import uuid
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from llm.client import LLMClient

class BaseTemplatesGenerator:
    """
    基于模板生成文档的类，用于根据预定义模板生成文档
    """
    
    def __init__(self, model_type: Optional[str] = None):
        """
        初始化模板生成器
        
        Args:
            model_type: 模型类型，"local"或"third_party"，默认使用环境变量中配置的类型
        """
        self.llm_client = LLMClient(model_type=model_type or "third_party")
        self.active_generations = {}
        
        # 创建必要的目录
        Path("downloads").mkdir(exist_ok=True)
        Path("templates").mkdir(exist_ok=True)
    
    def list_templates(self) -> List[str]:
        """
        列出所有可用的模板
        
        Returns:
            模板名称列表
        """
        templates_dir = Path("templates")
        templates = [f.stem for f in templates_dir.glob("*.docx")]
        return templates
    
    def generate_from_template(self, template_name: str, data: Dict[str, Any]) -> str:
        """
        根据模板和数据生成文档
        
        Args:
            template_name: 模板名称
            data: 用于填充模板的数据
            
        Returns:
            生成的文档路径
        """
        if not template_name:
            raise ValueError("模板名称不能为空")
        
        template_path = os.path.join("templates", f"{template_name}.docx")
        if not os.path.exists(template_path):
            raise ValueError(f"模板 '{template_name}' 不存在")
        
        try:
            # 打开模板文档
            doc = Document(template_path)
            
            # 填充模板内容
            self._fill_template(doc, data)
            
            # 生成随机文件名，避免冲突
            filename = f"{uuid.uuid4().hex}.docx"
            file_path = os.path.join("downloads", filename)
            doc.save(file_path)
            
            return file_path
        
        except Exception as e:
            raise Exception(f"生成文档时出错: {str(e)}")
    
    def _fill_template(self, doc: Document, data: Dict[str, Any]) -> None:
        """
        填充文档模板
        
        Args:
            doc: 文档对象
            data: 用于填充的数据
        """
        # 遍历所有段落
        for paragraph in doc.paragraphs:
            self._replace_text_in_paragraph(paragraph, data)
        
        # 遍历所有表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_text_in_paragraph(paragraph, data)
    
    def _replace_text_in_paragraph(self, paragraph, data: Dict[str, Any]) -> None:
        """
        在段落中替换占位符，支持占位符被拆分到多个run中的情况。
        这是一个复杂的操作，因为占位符可能横跨多个具有不同格式的run。
        """
        for key, value in data.items():
            # 同时兼容【】和{{}}两种格式
            placeholders = [f"【{key}】", f"{{{{{key}}}}}"]
            for p_text in placeholders:
                # 使用while循环处理段落中多次出现的同一占位符
                while p_text in paragraph.text:
                    runs = paragraph.runs
                    full_text = "".join(r.text for r in runs)

                    if p_text not in full_text:
                        break  # 如果在文本中找不到，退出循环

                    # --- 定位占位符所在的 runs ---
                    start_char_index = full_text.find(p_text)
                    end_char_index = start_char_index + len(p_text)
                    
                    run_cursor = 0
                    start_run, end_run = None, None
                    start_run_idx, end_run_idx = -1, -1
                    start_offset, end_offset = -1, -1

                    for i, run in enumerate(runs):
                        run_len = len(run.text)
                        if start_run is None and run_cursor + run_len > start_char_index:
                            start_run = run
                            start_run_idx = i
                            start_offset = start_char_index - run_cursor
                        if end_run is None and run_cursor + run_len >= end_char_index:
                            end_run = run
                            end_run_idx = i
                            end_offset = end_char_index - run_cursor
                            break
                        run_cursor += run_len
                    
                    # --- 执行替换 ---
                    if start_run is None or end_run is None:
                        # 如果没有找到起始或结束run，可能存在问题，跳出循环
                        break

                    if start_run_idx == end_run_idx:
                        # 占位符在单个run内
                        start_run.text = start_run.text[:start_offset] + str(value) + end_run.text[end_offset:]
                    else:
                        # 占位符横跨多个run
                        # 1. 修改起始run
                        start_run.text = start_run.text[:start_offset] + str(value)
                        # 2. 修改结束run
                        end_run.text = end_run.text[end_offset:]
                        # 3. 清空中间的runs
                        for i in range(start_run_idx + 1, end_run_idx):
                            runs[i].text = ""
    
    def create_template(self, template_name: str, structure: Dict[str, Any]) -> str:
        """
        创建新的文档模板
        
        Args:
            template_name: 模板名称
            structure: 模板结构定义
            
        Returns:
            创建的模板路径
        """
        if not template_name:
            raise ValueError("模板名称不能为空")
        
        template_path = os.path.join("templates", f"{template_name}.docx")
        
        try:
            # 创建新文档
            doc = Document()
            
            # 设置文档基本属性
            core_properties = doc.core_properties
            core_properties.title = template_name
            core_properties.language = "zh-CN"
            
            # 设置页面边距和纸张大小
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)
                section.page_height = Inches(11.69)  # A4纸高度
                section.page_width = Inches(8.27)    # A4纸宽度
            
            # 根据结构创建模板内容
            self._build_template_structure(doc, structure)
            
            # 添加页脚
            self._add_footer(doc)
            
            # 保存模板
            doc.save(template_path)
            
            return template_path
        
        except Exception as e:
            raise Exception(f"创建模板时出错: {str(e)}")
    
    def _build_template_structure(self, doc: Document, structure: Dict[str, Any]) -> None:
        """
        构建模板结构
        
        Args:
            doc: 文档对象
            structure: 模板结构定义
        """
        # 添加标题
        if "title" in structure:
            title = doc.add_heading(structure["title"], level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加章节
        if "sections" in structure:
            for section in structure["sections"]:
                if "heading" in section:
                    doc.add_heading(section["heading"], level=section.get("level", 1))
                
                if "content" in section:
                    for content_item in section["content"]:
                        if content_item.get("type") == "paragraph":
                            doc.add_paragraph(content_item.get("text", ""))
                        elif content_item.get("type") == "placeholder":
                            placeholder_text = f"【{content_item.get('name', 'placeholder')}】"
                            doc.add_paragraph(placeholder_text)
                        elif content_item.get("type") == "table" and "rows" in content_item and "cols" in content_item:
                            self._add_table(doc, content_item)
    
    def _add_table(self, doc: Document, table_def: Dict[str, Any]) -> None:
        """
        添加表格到文档
        
        Args:
            doc: 文档对象
            table_def: 表格定义        
        """
        rows = table_def.get("rows", 1)
        cols = table_def.get("cols", 1)
        
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        # 填充表头
        if "headers" in table_def and len(table_def["headers"]) > 0:
            for i, header in enumerate(table_def["headers"]):
                if i < cols:
                    table.cell(0, i).text = header
        
        # 填充占位符
        if "placeholders" in table_def:
            for row_idx, row_data in enumerate(table_def["placeholders"]):
                if row_idx + 1 < rows:  # +1 因为第一行可能是表头
                    for col_idx, placeholder in enumerate(row_data):
                        if col_idx < cols:
                            placeholder_text = f"【{placeholder}】"
                            table.cell(row_idx + 1, col_idx).text = placeholder_text
    
    def _add_footer(self, doc: Document) -> None:
        """
        添加页脚到文档
        
        Args:
            doc: 文档对象
        """
        for section in doc.sections:
            footer = section.footer
            p = footer.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.text = "第 "
            run = p.add_run()
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'PAGE'
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            run._element.append(fldChar)
            run._element.append(instrText)
            run._element.append(fldChar2)
            p.add_run(" 页")
    
    def generate_template_with_ai(self, template_name: str, description: str) -> str:
        """
        使用AI生成模板结构
        
        Args:
            template_name: 模板名称
            description: 模板描述
            
        Returns:
            创建的模板路径
        """
        if not template_name or not description:
            raise ValueError("模板名称和描述不能为空")
        
        try:
            # 构建提示，要求AI生成模板结构
            prompt = f"请根据以下描述生成一个文档模板的结构定义，以JSON格式返回：\n\n描述：{description}\n\n"
            prompt += "结构应包含标题、章节、内容块（段落、占位符、表格等）。请确保JSON格式正确，可以直接被Python解析。"
            
            messages = [
                {"role": "system", "content": "你是一个专业的文档模板设计助手，擅长根据描述创建结构化的文档模板。"},
                {"role": "user", "content": prompt}
            ]
            
            # 调用LLM生成模板结构
            response = self.llm_client.chat(messages=messages)
            structure_text = response.get("choices", [{}])[0].get("message", {}).get("content", "")
            
            if not structure_text:
                raise ValueError("模型返回内容为空")
            
            # 提取JSON部分
            import json
            import re
            
            # 尝试找到JSON部分
            json_match = re.search(r'```json\s*([\s\S]*?)\s*```', structure_text)
            if json_match:
                structure_json = json_match.group(1)
            else:
                structure_json = structure_text
            
            # 解析JSON
            try:
                structure = json.loads(structure_json)
            except json.JSONDecodeError:
                # 如果解析失败，尝试清理和修复JSON
                cleaned_json = re.sub(r'[\n\r\t]', '', structure_json)
                cleaned_json = re.sub(r',\s*}', '}', cleaned_json)
                structure = json.loads(cleaned_json)
            
            # 创建模板
            return self.create_template(template_name, structure)
        
        except Exception as e:
            raise Exception(f"使用AI生成模板时出错: {str(e)}")
    
    def cleanup_old_files(self, max_age_hours: int = 24) -> None:
        """
        清理长时间未使用的生成文件
        
        Args:
            max_age_hours: 文件最大保留时间（小时）
        """
        current_time = time.time()
        downloads_dir = Path("downloads")
        
        for file_path in downloads_dir.glob("*.docx"):
            file_age = current_time - os.path.getmtime(file_path)
            if file_age > (max_age_hours * 3600):  # 转换为秒
                try:
                    os.remove(file_path)
                except Exception as e:
                    print(f"删除文件 {file_path} 时出错: {str(e)}")
    
    def extract_variables_from_template(self, template_path: str) -> List[Dict[str, str]]:
        """
        从Word模板中提取所有变量
        
        Args:
            template_path: 模板文件路径
            
        Returns:
            变量列表，每个变量是一个包含name和description的字典
        """
        import re
        
        if not os.path.exists(template_path):
            raise ValueError(f"模板文件 '{template_path}' 不存在")
        
        try:
            # 打开模板文档
            doc = Document(template_path)
            
            # 存储找到的变量
            variables = []
            found_vars = set()
            
            # 正则表达式匹配【变量名】或{{变量名}}格式
            var_pattern = re.compile(r'【([^】]+?)】|\{\{([^}]+?)\}\}')
            
            # 从段落中提取变量
            for paragraph in doc.paragraphs:
                matches = var_pattern.findall(paragraph.text)
                for match_tuple in matches:
                    # findall会返回元组，(中文匹配, 英文匹配)，取非空的一个
                    var_name = next((v for v in match_tuple if v), None)
                    if var_name:
                        var_name = var_name.strip()
                        if var_name and var_name not in found_vars:
                            found_vars.add(var_name)
                            variables.append({
                                "name": var_name,
                                "description": f"{var_name}的值",
                                "default_value": ""
                            })
            
            # 从表格中提取变量
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            matches = var_pattern.findall(paragraph.text)
                            for match_tuple in matches:
                                var_name = next((v for v in match_tuple if v), None)
                                if var_name:
                                    var_name = var_name.strip()
                                    if var_name and var_name not in found_vars:
                                        found_vars.add(var_name)
                                        variables.append({
                                            "name": var_name,
                                            "description": f"{var_name}的值",
                                            "default_value": ""
                                        })
            
            return variables
        
        except Exception as e:
            raise Exception(f"提取模板变量时出错: {str(e)}")
    
    def register_uploaded_template(self, template_path: str) -> Dict[str, Any]:
        """
        注册上传的模板，提取变量并保存元数据
        
        Args:
            template_path: 模板文件路径
            
        Returns:
            包含变量和元数据的字典
        """
        try:
            # 提取变量
            variables = self.extract_variables_from_template(template_path)
            
            # 构建元数据
            template_name = os.path.basename(template_path).split('.')[0]
            metadata = {
                "name": template_name,
                "type": "uploaded",
                "variables": variables,
                "created_at": time.strftime("%Y-%m-%d %H:%M:%S")
            }
            
            # 可以将元数据保存到文件中，但这里我们直接返回
            return metadata
        
        except Exception as e:
            raise Exception(f"注册模板时出错: {str(e)}")

    def get_template_content(self, template_path: str) -> str:
        """
        获取模板文件的文本内容
        
        Args:
            template_path: 模板文件路径
            
        Returns:
            模板的文本内容
        """
        try:
            if not os.path.exists(template_path):
                raise ValueError(f"模板文件 '{template_path}' 不存在")
            
            # 打开模板文档
            doc = Document(template_path)
            
            # 收集所有文本内容
            content_parts = []
            
            # 从段落中提取文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)
            
            # 从表格中提取文本
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        content_parts.append(' | '.join(row_texts))
            
            # 合并所有内容
            return '\n'.join(content_parts)
            
        except Exception as e:
            raise Exception(f"获取模板内容时出错: {str(e)}")

if __name__ == "__main__":
    generator = BaseTemplatesGenerator(model_type="third_party")
    generator.generate_template_with_ai("test", "这是一个测试模板")
    print(generator.list_templates())
